from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
GENERATED = ROOT / "data" / "generated"
REPORTS = ROOT / "reports"


def load_json(name: str) -> dict:
    return json.loads((GENERATED / name).read_text(encoding="utf-8"))


def set_normal_style(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.25
    style.paragraph_format.space_after = Pt(6)


def add_title(document: Document, title: str, subtitle: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(21, 94, 117)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(71, 85, 105)


def add_kv_table(document: Document, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Cm(4.2)
    table.columns[1].width = Cm(11.8)
    hdr = table.rows[0].cells
    hdr[0].text = "项目"
    hdr[1].text = "内容"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_picture_if_exists(document: Document, path: Path, caption: str) -> None:
    if not path.exists():
        return
    document.add_picture(str(path), width=Inches(5.8))
    p = document.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(71, 85, 105)


def build_report() -> Path:
    REPORTS.mkdir(parents=True, exist_ok=True)
    data_report = load_json("data_report.json")
    evaluation = load_json("evaluation_report.json")
    stats = load_json("stats_report.json")

    triples_path = GENERATED / "knowledge_triples.csv"
    triple_count = max(len(triples_path.read_text(encoding="utf-8-sig").splitlines()) - 1, 0)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    set_normal_style(doc)

    add_title(doc, "基于 CBLUE 中文基准的医学文本实体关系抽取与知识图谱构建", "终期结题报告")
    add_kv_table(
        doc,
        [
            ("项目名称", "基于 CBLUE 中文基准的医学文本实体关系抽取与知识图谱构建"),
            ("数据来源", "CBLUE 中 CMeEE、CMeIE、KUAKE-QIC 子任务"),
            ("核心方法", "词典基线、BERT 训练入口、关系三元组抽取、Mann-Whitney U 检验、Neo4j 导出"),
            ("主要产物", "代码仓库、Streamlit Demo、结题报告、汇报 PPT、Neo4j Cypher、实验图表"),
        ],
    )

    doc.add_heading("1. 研究背景与问题定义", level=1)
    doc.add_paragraph(
        "中文医学文本包含大量专业实体、缩写、同义表达和口语化主诉。"
        "在病历结构化、门诊导诊和医学知识整理中，实体识别和关系抽取是后续推理与检索的基础。"
        "本项目选用 CBLUE 中文生物医学语言理解基准中的 CMeEE、CMeIE 和 KUAKE-QIC 数据，"
        "围绕实体抽取、关系抽取、意图分类结果汇总和知识图谱构建开展实验。"
    )
    add_bullets(
        doc,
        [
            "CMeEE：从医学文本中识别疾病、症状、身体部位等实体。",
            "CMeIE：从医学文本中抽取主体、关系、客体三元组。",
            "KUAKE-QIC：对医学查询进行意图分类，并汇总已有训练日志中的模型结果。",
        ],
    )

    doc.add_heading("2. 技术路线", level=1)
    doc.add_paragraph(
        "项目流程从原始医学文本开始，先进行数据读取、清洗和统计，再分别进入描述统计、实体识别、关系抽取和知识图谱导出。"
        "实体侧采用高频词典作为可解释基线，并补充 BERT token classification 训练脚本；关系侧从 CMeIE 中抽取三元组，"
        "写出 CSV 和 Neo4j Cypher。统计部分使用 Mann-Whitney U 检验比较多次实验分数差异。"
    )
    add_bullets(
        doc,
        [
            "数据层：读取 CBLUE 子任务数据，统计样本数量、文本长度和标签分布。",
            "模型层：实现 CMeEE 词典基线，保留 CMeEE BERT 训练入口。",
            "评估层：计算 Precision、Recall、F1、Accuracy 和 p-value。",
            "应用层：导出 Neo4j 导入脚本，并通过 Streamlit 页面展示结果。",
        ],
    )

    doc.add_heading("3. 数据集与工程结构", level=1)
    datasets = data_report["datasets"]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["文件", "样本数", "平均长度", "最大长度"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in datasets:
        cells = table.add_row().cells
        cells[0].text = str(row["file"])
        cells[1].text = str(row["samples"])
        cells[2].text = str(row["avg_text_len"])
        cells[3].text = str(row["max_text_len"])
    doc.add_paragraph(
        "仓库按模块拆分为 cblue_project、scripts、ui、docs、tests、reports 等目录。"
        "其中 cblue_project 放核心逻辑，scripts 放构建和训练入口，ui 放 Streamlit 页面，tests 放单元测试。"
    )

    doc.add_heading("4. 实验设计与结果", level=1)
    cmeee = evaluation["cmeee_dictionary_baseline"]
    qic = evaluation["qic_chinese_bert_wwm_ext"]
    mwu = stats["mann_whitney_u"]
    add_kv_table(
        doc,
        [
            ("CMeEE 词典基线 Precision", str(cmeee["precision"])),
            ("CMeEE 词典基线 Recall", str(cmeee["recall"])),
            ("CMeEE 词典基线 F1", str(cmeee["f1"])),
            ("KUAKE-QIC 最佳开发集 Accuracy", str(qic["best_dev_accuracy"])),
            ("Mann-Whitney U p-value", str(mwu["p_value"])),
            ("CMeIE 三元组规模", f"{triple_count} 条"),
        ],
    )
    doc.add_paragraph(
        "CMeEE 词典基线结果说明，单纯依靠高频实体匹配可以形成可解释的低成本基线，但召回和精度都受到同义表达、"
        "实体边界和短实体误匹配影响，因此 F1 仍然偏低。后续训练脚本已经补齐，适合继续完成 BERT 微调。"
    )
    add_picture_if_exists(doc, GENERATED / "figures" / "metrics_summary.png", "图 1：当前实验指标汇总")
    add_picture_if_exists(doc, GENERATED / "figures" / "score_comparison.png", "图 2：多次实验分数对比")

    doc.add_heading("5. 知识图谱构建", level=1)
    doc.add_paragraph(
        f"项目从 CMeIE 数据中抽取并去重得到 {triple_count} 条三元组，字段包括 subject、subject_type、predicate、object、object_type 和 source_text。"
        "这些结果写入 knowledge_triples.csv，同时生成 neo4j_import.cypher，便于在 Neo4j Browser 或 cypher-shell 中导入。"
    )
    add_bullets(
        doc,
        [
            "节点：统一使用 Entity 标签，并保留实体类型。",
            "关系：统一使用 RELATED_TO，predicate 属性记录原始关系类型。",
            "来源：source 属性保留截断后的原文片段，方便回溯。",
        ],
    )

    doc.add_heading("6. 可复现性与测试", level=1)
    doc.add_paragraph(
        "仓库没有提交原始 CBLUE 数据集和预训练模型，原因是文件体积较大。docs/reproducibility.md 中写明了所需目录结构、"
        "训练日志格式和运行命令。Dockerfile 已补充，单元测试覆盖了指标计算、日志解析和 Cypher 导出等核心逻辑。"
    )
    add_bullets(
        doc,
        [
            "生成产物：python scripts/build_artifacts.py",
            "启动页面：streamlit run ui/triage_simple_app.py",
            "训练入口：python scripts/train_cmeee_bert.py",
            "测试命令：python -m unittest discover -s tests",
        ],
    )

    doc.add_heading("7. 不足与后续工作", level=1)
    add_bullets(
        doc,
        [
            "CMeEE 当前主结果仍是词典基线，后续应完整运行 BERT 微调并记录训练条件。",
            "KUAKE-QIC 当前引用本地日志结果，后续应补充独立复现实验和更多随机种子。",
            "消融实验还可以继续补充，例如去掉实体去重、去掉实体融合或更换关系抽取规则。",
            "知识图谱可以继续扩大三元组规模，并补充查询模板和可视化截图。",
        ],
    )

    doc.add_heading("8. 总结", level=1)
    doc.add_paragraph(
        "本项目完成了 CBLUE 数据读取、CMeEE 基线评估、KUAKE-QIC 日志结果汇总、Mann-Whitney U 检验、"
        "CMeIE 三元组导出、Neo4j 脚本生成和 Streamlit 页面展示。终期阶段补充了结题报告、汇报 PPT、"
        "单元测试、Dockerfile、复现说明和图表产物。项目已经形成较完整的中文医学文本处理实验链路，"
        "后续重点是扩大模型训练和对照实验规模。"
    )

    output = REPORTS / "CBLUE_结题报告.docx"
    doc.save(output)
    return output


if __name__ == "__main__":
    print(build_report())
